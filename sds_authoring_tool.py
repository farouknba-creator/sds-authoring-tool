import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
import json
import os
import sys
import re
from docx import Document
import pdfplumber

# -------------------------------------------------------------------
# Helpers for PyInstaller bundled files
# -------------------------------------------------------------------
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def load_json(file_name, default=None):
    """Safely load a JSON file, returning default if missing or corrupted."""
    path = resource_path(file_name)
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return default

# -------------------------------------------------------------------
# Extract sections from PDF (improved heading detection)
# -------------------------------------------------------------------
def extract_sections_from_pdf(pdf_path):
    """Extract sections from a GHS‑compliant SDS PDF. Handles:
       - SECTION 1: Identification
       - Section 1. Identification
       - 1. Identification
       - 1. HAZARDS IDENTIFICATION
       and similar variations.
    """
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        return [f"Error extracting PDF: {e}"]

    if not text.strip():
        return ["[No extractable text – PDF may be scanned.]"]

    # Normalise line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Pattern to match GHS section headings, capturing the full heading line
    # Matches: SECTION 1: Identification / Section 1. Identification / 1. Identification / etc.
    heading_pattern = re.compile(
        r'^(?:SECTION\s+)?(\d{1,2})\s*[\.\:\)\-]?\s*(.*?)\s*$',
        re.IGNORECASE | re.MULTILINE
    )

    # Find all matches with their positions
    matches = list(re.finditer(heading_pattern, text))
    if not matches:
        # No headings found – return the whole text as one block
        return [text.strip()]

    sections = []
    # Split at each heading; discard any preamble before the first heading
    for i, match in enumerate(matches):
        section_num = match.group(1)
        title = match.group(2).strip()
        start = match.start()
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        body = text[start:end].strip()
        # Create a clean label for display (not saved, just for the listbox)
        label = f"Section {section_num} – {title}" if title else f"Section {section_num}"
        # The section content is the whole block from the heading onward
        sections.append(body)

    return sections

# -------------------------------------------------------------------
# Main Application
# -------------------------------------------------------------------
class SDSAuthoringTool(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("SDS Authoring Tool")
        self.geometry("1100x750")
        self.resizable(True, True)

        # Data
        self.pdf_path = None
        self.source_sections = []              # full text of each section
        self.template_fields = []              # from template_fields.json
        self.mappings = {}                     # label -> section_text
        self.identities = []
        self.current_identity = None

        # Load config files
        self.load_configs()
        self.create_widgets()

    def load_configs(self):
        field_cfg = load_json("template_fields.json", default={"fields": [], "identity_placeholders": {}})
        self.template_fields = field_cfg.get("fields", [])
        self.identities = load_json("identities.json", default=[])

    def create_widgets(self):
        # Top bar
        top_frame = tk.Frame(self)
        top_frame.pack(fill="x", padx=10, pady=5)

        tk.Button(top_frame, text="Open PDF", command=self.load_pdf).pack(side="left", padx=5)
        tk.Label(top_frame, text="Identity set:").pack(side="left", padx=(20,5))
        self.identity_var = tk.StringVar()
        self.identity_combo = ttk.Combobox(top_frame, textvariable=self.identity_var,
                                           state="readonly", width=30)
        self.identity_combo.pack(side="left", padx=5)
        self.identity_combo.bind("<<ComboboxSelected>>", self.on_identity_selected)
        self.populate_identity_list()

        # Main area: two listboxes + map button
        listbox_frame = tk.Frame(self)
        listbox_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Left: Source sections
        left_frame = tk.LabelFrame(listbox_frame, text="Source PDF Sections (select one)")
        left_frame.pack(side="left", fill="both", expand=True, padx=5)
        self.source_listbox = tk.Listbox(left_frame, selectmode="single", width=40, height=15)
        self.source_listbox.pack(fill="both", expand=True, padx=5, pady=5)

        # Middle: Map button
        mid_frame = tk.Frame(listbox_frame)
        mid_frame.pack(side="left", fill="y", padx=10)
        tk.Label(mid_frame, text="↓").pack(pady=5)
        tk.Button(mid_frame, text="Map\nSelected", command=self.map_selected, width=10, height=3).pack(pady=20)
        tk.Button(mid_frame, text="Clear\nMappings", command=self.clear_mappings, width=10, height=3).pack(pady=10)

        # Right: Template fields
        right_frame = tk.LabelFrame(listbox_frame, text="Template Fields (select one)")
        right_frame.pack(side="left", fill="both", expand=True, padx=5)
        self.fields_listbox = tk.Listbox(right_frame, selectmode="single", width=40, height=15)
        self.fields_listbox.pack(fill="both", expand=True, padx=5, pady=5)

        # Populate fields list
        self.refresh_fields_display()

        # Raw text preview (collapsible)
        preview_frame = tk.LabelFrame(self, text="Extracted Raw Text (for diagnostics)")
        preview_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.preview_text = scrolledtext.ScrolledText(preview_frame, height=8, width=80, state="disabled")
        self.preview_text.pack(fill="both", expand=True, padx=5, pady=5)

        # Bottom bar
        bottom_frame = tk.Frame(self)
        bottom_frame.pack(fill="x", padx=10, pady=5)
        tk.Button(bottom_frame, text="Export DOCX", command=self.export_docx, height=2).pack(side="right", padx=5)
        self.status = tk.Label(self, text="Ready", bd=1, relief="sunken", anchor="w")
        self.status.pack(side="bottom", fill="x")

    def populate_identity_list(self):
        names = [ident["name"] for ident in self.identities]
        self.identity_combo["values"] = names
        if names:
            self.identity_combo.current(0)
            self.current_identity = self.identities[0]

    def on_identity_selected(self, event):
        idx = self.identity_combo.current()
        if 0 <= idx < len(self.identities):
            self.current_identity = self.identities[idx]
            self.status.config(text=f"Selected identity: {self.current_identity['name']}")

    def load_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if not path:
            return
        self.pdf_path = path
        self.status.config(text="Extracting sections from PDF...")
        self.update()
        self.source_sections = extract_sections_from_pdf(path)
        self.source_listbox.delete(0, tk.END)
        for sec in self.source_sections:
            # Show a short preview for the listbox (first 120 chars)
            preview = sec[:120].replace("\n", " ") + ("..." if len(sec) > 120 else "")
            self.source_listbox.insert(tk.END, preview)
        # Show full raw text in the preview area
        self.show_raw_text()
        self.clear_mappings()
        self.status.config(text=f"Loaded {len(self.source_sections)} sections from PDF.")

    def show_raw_text(self):
        """Display the full extracted text (concatenation of all sections) in the preview pane."""
        raw = "\n\n--- Section break ---\n\n".join(self.source_sections)
        self.preview_text.config(state="normal")
        self.preview_text.delete(1.0, tk.END)
        self.preview_text.insert(tk.END, raw)
        self.preview_text.config(state="disabled")

    def refresh_fields_display(self):
        """Update the fields listbox to show current mapping status."""
        self.fields_listbox.delete(0, tk.END)
        for field in self.template_fields:
            label = field["label"]
            if label in self.mappings:
                mapped_text = self.mappings[label]
                short = mapped_text[:50].replace("\n", " ") + "..."
                self.fields_listbox.insert(tk.END, f"[✓] {label} – {short}")
            else:
                self.fields_listbox.insert(tk.END, f"[  ] {label}")

    def map_selected(self):
        """Map the selected source section to the selected template field."""
        source_sel = self.source_listbox.curselection()
        target_sel = self.fields_listbox.curselection()

        # Diagnostic: show raw selection indices in status bar
        self.status.config(text=f"Source sel: {source_sel}, Target sel: {target_sel}")
        self.update()

        if not source_sel:
            messagebox.showwarning("Mapping", "Please select a source section first.")
            return
        if not target_sel:
            messagebox.showwarning("Mapping", "Please select a template field first.")
            return

        src_idx = source_sel[0]
        tgt_idx = target_sel[0]

        if src_idx >= len(self.source_sections) or tgt_idx >= len(self.template_fields):
            messagebox.showerror("Error", "Selected index out of range.")
            return

        section_text = self.source_sections[src_idx]
        field = self.template_fields[tgt_idx]
        label = field["label"]

        self.mappings[label] = section_text
        self.refresh_fields_display()
        self.status.config(text=f"Mapped section {src_idx+1} → {label}")

    def clear_mappings(self):
        self.mappings = {}
        self.refresh_fields_display()
        self.status.config(text="Mappings cleared.")

    def export_docx(self):
        if not self.pdf_path:
            messagebox.showerror("Error", "No PDF loaded.")
            return
        output_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                   filetypes=[("Word Document", "*.docx")])
        if not output_path:
            return

        template_path = resource_path("template.docx")
        if not os.path.exists(template_path):
            messagebox.showerror("Error", "template.docx not found in app folder.")
            return

        doc = Document(template_path)
        replacements = {}

        # Identity placeholders
        if self.current_identity:
            id_cfg = load_json("template_fields.json", default={}).get("identity_placeholders", {})
            for key, placeholder in id_cfg.items():
                value = self.current_identity.get(key, "")
                replacements[placeholder] = value

        # Section placeholders from mappings
        for field in self.template_fields:
            placeholder = "{{" + field["placeholder"] + "}}"
            label = field["label"]
            replacements[placeholder] = self.mappings.get(label, "")

        # Replace in paragraphs
        for para in doc.paragraphs:
            for ph, val in replacements.items():
                if ph in para.text:
                    if para.text.strip() == ph:
                        para.clear()
                        para.add_run(val)
                    else:
                        para.text = para.text.replace(ph, val)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for ph, val in replacements.items():
                            if ph in para.text:
                                if para.text.strip() == ph:
                                    para.clear()
                                    para.add_run(val)
                                else:
                                    para.text = para.text.replace(ph, val)

        doc.save(output_path)
        messagebox.showinfo("Success", f"Document saved to:\n{output_path}")
        self.status.config(text=f"Exported to {output_path}")

if __name__ == "__main__":
    app = SDSAuthoringTool()
    app.mainloop()
